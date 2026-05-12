import json
import os
from google import genai
from docx import Document
from docxtpl import DocxTemplate

from config import COMPANY_NAME, JOB_TITLE, OUTPUT_FOLDER_COVERLETTER, RESUME_SOURCE, CV_SOURCE, JOB_DESCRIPTION

client = genai.Client()

# job_url = "https://www.isnetworld.com/en/careers/5977058004?gh_src=e5150d284us#application"
# page = requests.get(job_url)
# soup = BeautifulSoup(page.content, "html.parser")
job_posting_text = JOB_DESCRIPTION

source_doc = Document(RESUME_SOURCE)
full_resume_text = "\n".join([p.text for p in source_doc.paragraphs])

reader_cv = Document(CV_SOURCE)
cv_text = "\n".join([p.text for p in reader_cv.paragraphs])

background_info = f"RESUME DATA:\n{full_resume_text}\n\nCV DATA:\n{cv_text}"

prompt=f"""
I am applying for a role at {COMPANY_NAME}
Job requirements: {job_posting_text}
Here is my full professional background: {background_info}

Rewrite my cover letter to align with the job's key requirements.

Write a professional cover letter that is exactly 3-4 paragraphs long and fits on one page.

STRICT COVER LETTER INSTRUCTIONS:
1. CUSTOMIZATION: Mention the {JOB_TITLE} role at {COMPANY_NAME} in the first paragraph.
2. PROBLEM-SOLVER TONE: Focus on how my background in 'Root Cause Identification' and 'SQL querying' (from my OSU and NASA research) makes me a perfect fit for the DevSupport team.
3. THE "WHY": Connect my passion for mission-driven engineering to {COMPANY_NAME}'s goals.
4. BREVITY: Each paragraph should be no more than 4-5 sentences.
5. NO HALLUCINATION: Only mention skills found in my resume/CV. If the job asks for ASP.NET and I don't have it, focus on my strong C# and web development foundation instead.
6. SALUTATION: Start the 'introduction' section with "Dear {COMPANY_NAME} Hiring Manager," or "Dear {COMPANY_NAME} Hiring Team,".
7. ONE-PAGE GOAL: Use concise, professional language. Avoid filler words like "responsible for" or "duties included."

Return ONLY a JSON object with these keys: 
"companyName", "introduction", "body_para_1", "body_para_2", "closing"

ATS OPTIMIZATION RULES:
1. KEYWORD DENSITY: Naturally integrate phrases like "Root Cause Identification", "T-SQL", and "data-ingestion pipeline".
2. BRANDING: Mention {COMPANY_NAME} and the {JOB_TITLE} role in the first 2 sentences.
3. ACTION VERBS: Start sentences with verbs like "Engineered", "Migrated", or "Troubleshot".
4. NO FORMATTING: Return plain text without Markdown (no asterisks or bolding) so the Python script can handle the Word formatting.
"""

response = client.models.generate_content(
    model="gemini-2.5-flash",
    contents=prompt,
    config={'response_mime_type': 'application/json'}
)

ai_results = json.loads(response.text)

template_doc = DocxTemplate('CoverLetter-TEMPLATE.docx')

replacements = {
    "COMPANYNAME": ai_results.get("company_name"),
    "INTRODUCTION": ai_results.get("introduction", ""),
    "BODY_PARA_1": ai_results.get("body_para_1", ""),
    "BODY_PARA_2": ai_results.get("body_para_2", ""),
    "CLOSING": ai_results.get("closing", "")
}

template_doc.render(replacements)

filename = f"{COMPANY_NAME}_IsabellaMann_CoverLetter.docx"
save_path = os.path.join(OUTPUT_FOLDER_COVERLETTER, filename)

template_doc.save(save_path)
print(f"Done! saved to: {save_path}")