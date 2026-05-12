import json
import os
from google import genai
from docx import Document
from docxtpl import DocxTemplate

from config import COMPANY_NAME, JOB_TITLE, OUTPUT_FOLDER_RESUME, RESUME_SOURCE, CV_SOURCE, JOB_DESCRIPTION


client = genai.Client()

# job_url = "https://www.isnetworld.com/en/careers/5977058004?gh_src=e5150d284us#application"
# page = requests.get(job_url)
# soup = BeautifulSoup(page.content, "html.parser")

source_doc = Document(RESUME_SOURCE)
full_resume_text = "\n".join([p.text for p in source_doc.paragraphs])

reader_cv = Document(CV_SOURCE)
cv_text = "\n".join([p.text for p in reader_cv.paragraphs])

background_info = f"RESUME DATA:\n{full_resume_text}\n\nCV DATA:\n{cv_text}"

prompt=f"""
I am applying for a role at {COMPANY_NAME}
Job requirements: {JOB_DESCRIPTION}
Job title: {JOB_TITLE}
Here is my full professional background: {background_info}

Rewrite my 'Professional Summary', 'Experience', 'Projects', and 'Skills' sections
to align with the job's key requirements.

SKILLS INTEGRATION & FORMATTING:
1. INTEGRATION: Weave technical tools (e.g., SQL, Jira, MongoDB) directly into the experience bullets to show context.
2. SKILLS SECTION: Keep the exact format from the original resume: 
   Languages: [List]
   Frameworks & Libraries: [List]
   Tools and Environments: [List]
3. CAP SKILLS: Limit each category in the Skills section to the top 6-9 most relevant items.
4. CAP EXPERIENCES: Limit the number of experiences in the Experience section to the top 4 most relevant items.
5. ONE-PAGE GOAL: Use concise, professional language. Avoid filler words like "responsible for" or "duties included."
6. PROJECT LIMIT: Limit each project to 2-3 concise bullet points.

CRITICAL ATS (Applicant Tracking System) & BACKGROUND INSTRUCTIONS:
1. TRUTH-ONLY RULE: Use your analysis of my provided resume to determine my actual skills. 
2. NO HALLUCINATION: If a requirement in the job posting (e.g., "ASP.NET") is NOT found in my resume, DO NOT add it to my experience or skills. 
3. ALTERNATIVE PHRASING: If I have a related skill but not the exact one (e.g., I have "React" but the job asks for "Angular"), highlight my "JavaScript frameworks" and "fast learning" instead of claiming the missing skill.
4. SYNONYM MAPPING: Only use exact phrases from the job description (like "T-SQL" or "Root Cause Identification") if I have already listed the base skill (like "SQL" or "troubleshooting") in my resume.
5. CV UTILIZATION: Use the CV data to find specific metrics or technical details (like your work with ArcGIS or satellite data) that might be missing from the shorter resume but are relevant to this job

CRITICAL FORMATTING INSTRUCTIONS FOR JSON OUTPUT:
1. You must return the output ONLY as a JSON object using exactly these four keys: "summary", "experience", "projects", "skills".
2. The value for EVERY key must be a single, flat string of plain text. Do not use nested arrays or objects.
3. SPACING: For Experience and Projects, use the exact following layout for each item:
   [Job Title] | [Company], [Location]
   [Dates]
   • [Bullet point 1]
   • [Bullet point 2]
4. BULLET INDENTATION: For every bullet point in Experience and Projects, start the line with a tab space or four spaces followed by a bullet symbol (e.g., "    • "). This ensures the bullets don't sit flush against the left margin.
"""

response = client.models.generate_content(
    model="gemini-2.5-flash",
    contents=prompt,
    config={'response_mime_type': 'application/json'}
)

ai_results = json.loads(response.text)

template_doc = DocxTemplate('Resume-TEMPLATE.docx')

replacements = {
    "SUMMARY": ai_results.get("summary", ""),
    "EXPERIENCE": ai_results.get("experience", ""),
    "PROJECTS": ai_results.get("projects", ""),
    "SKILLS": ai_results.get("skills", "")
}

template_doc.render(replacements)

filename = f"{companyName}_IsabellaMann_Resume.docx"
save_path = os.path.join(output_folder, filename)

template_doc.save(save_path)
print(f"Done! saved to: {save_path}")